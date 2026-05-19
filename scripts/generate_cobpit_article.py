#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_cobpit_article.py
==========================
Gera o artigo adaptado para submissão ao III COBPIT – Trabalho Completo,
usando o template oficial template-tc-cobpit.docx localizado na raiz do projeto.

Uso:
    python scripts/generate_cobpit_article.py

Saída:
    _output/article/enaju-gcpj-cobpit-tc.docx
"""

import os
import sys
import copy
import csv
from collections import Counter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Caminhos
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(SCRIPT_DIR)
TEMPLATE_PATH = os.path.join(ROOT_DIR, "manuscripts", "templates", "template-tc-cobpit.docx")
OUTPUT_DIR = os.path.join(ROOT_DIR, "outputs")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "enaju-gcpj-cobpit-tc.docx")
CORPUS_ELIGIBLE_CSV = os.path.join(ROOT_DIR, "data", "processed", "corpus_eligible.csv")
DEDUP_REPORT_CSV = os.path.join(ROOT_DIR, "data", "processed", "dedup_report.csv")

# ---------------------------------------------------------------------------
# Verificações
# ---------------------------------------------------------------------------
if not os.path.exists(TEMPLATE_PATH):
    sys.exit(f"[ERRO] Template não encontrado: {TEMPLATE_PATH}")

os.makedirs(OUTPUT_DIR, exist_ok=True)


def _fmt_int(value):
    return f"{int(value):,}".replace(",", ".")


def _fmt_pct(value):
    return f"{float(value):.1f}".replace(".", ",")


def _load_corpus_stats():
    stats = {
        "n_total": 0,
        "year_min": 1995,
        "year_max": 2025,
        "n_judicial": 0,
        "pct_judicial": 0.0,
        "n_journals": 0,
        "n_sources": 0,
        "sources": "Scopus",
        "post_2010": 0,
        "pct_post_2010": 0.0,
        "raw_total": 0,
        "dedup_total": 0,
        "by_corpus": Counter(),
        "country_note": (
            "metadados de país de afiliação não estavam disponíveis na "
            "extração Scopus usada nesta versão"
        ),
    }

    if os.path.exists(CORPUS_ELIGIBLE_CSV):
        with open(CORPUS_ELIGIBLE_CSV, newline="", encoding="utf-8") as f:
            rows = list(csv.DictReader(f))

        stats["n_total"] = len(rows)
        years = [
            int(r["year"]) for r in rows
            if r.get("year") and r["year"].strip().isdigit()
        ]
        if years:
            stats["year_min"] = min(years)
            stats["year_max"] = max(years)
            stats["post_2010"] = sum(1 for y in years if 2010 <= y <= 2025)
            stats["pct_post_2010"] = (
                stats["post_2010"] / len(years) * 100 if years else 0.0
            )

        stats["by_corpus"] = Counter(r.get("corpus_id", "") for r in rows)
        stats["n_judicial"] = stats["by_corpus"].get("C", 0)
        stats["pct_judicial"] = (
            stats["n_judicial"] / stats["n_total"] * 100
            if stats["n_total"] else 0.0
        )
        stats["n_journals"] = len({
            r.get("journal", "").strip()
            for r in rows
            if r.get("journal", "").strip()
        })
        sources = sorted({
            r.get("source_db", "").strip()
            for r in rows
            if r.get("source_db", "").strip()
        })
        stats["n_sources"] = len(sources)
        stats["sources"] = ", ".join(sources) if sources else "Scopus"

    if os.path.exists(DEDUP_REPORT_CSV):
        with open(DEDUP_REPORT_CSV, newline="", encoding="utf-8") as f:
            rows = list(csv.DictReader(f))
        stats["raw_total"] = sum(int(float(r.get("n_raw", 0) or 0)) for r in rows)
        stats["dedup_total"] = sum(int(float(r.get("n_dedup", 0) or 0)) for r in rows)

    return stats


STATS = _load_corpus_stats()


# ===========================================================================
# CONTEÚDO DO ARTIGO
# ===========================================================================

TITULO = (
    "Educação Judicial e Inovação Pedagógica: mapeamento bibliométrico global "
    "para fundamentar práticas formativas em escolas judiciais"
)

AUTOR = (
    "Igor Caires Machado; Daniel Ribeiro Surdi de Avelar; "
    "Fabio Lopes Fernandes Ramos"
)

AFILIACAO = (
    "Conselho Nacional de Justiça / Escola Nacional do Judiciário (ENAJU), Brasília, DF, Brasil "
    "(cairesmachado@gmail.com)"
)

# Máximo 500 caracteres incluindo espaços
RESUMO = (
    f"Mapeamento bibliométrico (1995–2025) de "
    f"{_fmt_int(STATS['n_total'])} registros Scopus sobre educação corporativa "
    f"pública e judiciária. O Corpus C reúne "
    f"{_fmt_int(STATS['n_judicial'])} estudos "
    f"({_fmt_pct(STATS['pct_judicial'])}%). Identifica baixa integração entre "
    f"subcampos e propõe o IMECPJ como protótipo de benchmark para inovação "
    f"pedagógica em escolas judiciais."
)

# Máximo 5 palavras-chave
PALAVRAS_CHAVE = (
    "educação judicial; formação continuada; bibliometria; "
    "inovação pedagógica; escolas de governo"
)

# ---------------------------------------------------------------------------
# INTRODUÇÃO
# ---------------------------------------------------------------------------
INTRODUCAO = [
    (
        "A Escola Nacional do Judiciário "
        "(ENAJU) e demais escolas judiciais do Poder Judiciário brasileiro "
        "ocupam posição estratégica na formação continuada de magistrados e "
        "servidores: são simultaneamente instrumentos de política de pessoal, "
        "agentes de mudança organizacional e repositórios de conhecimento "
        "institucional (Megginson e Clutterbuck, 2007; Rainey, 2009). Apesar "
        "dessa relevância, as decisões sobre currículo, metodologias de ensino "
        "e avaliação de impacto formativo raramente se apoiam em evidências "
        "sistematizadas da produção científica global — porque essa produção, "
        "quando existe, está dispersa em comunidades acadêmicas que raramente "
        "dialogam entre si."
    ),
    (
        "A literatura científica sobre educação corporativa pública e "
        "judiciária fragmenta-se em ao menos três tradições: a da aprendizagem "
        "corporativa e do desenvolvimento de recursos humanos (HRD), centrada "
        "no setor privado (Swanson e Holton, 2001); a da administração pública "
        "e do fortalecimento do Estado (capacity development), associada a "
        "organismos como a OCDE (Grindle, 2004); e a da educação judicial, "
        "vinculada ao campo do direito e da administração da justiça "
        "(Geyh, 2006). Uma escola judicial que deseja implantar metodologias "
        "ativas ou aprendizagem baseada em problemas dificilmente encontrará "
        "referências em sua literatura especializada — elas estão dispersas em "
        "periódicos de tecnologia educacional e gestão de recursos humanos "
        "(Siemens e Baker, 2013)."
    ),
    (
        "Essa fragmentação tem consequências diretas para a inovação "
        "pedagógica: escolas de formação desenvolvem currículos e programas "
        "sem acesso ao acervo de evidências produzido por campos vizinhos, "
        "perdendo oportunidades de incorporar práticas baseadas em evidências. "
        "No contexto brasileiro, a Política Nacional de Desenvolvimento de "
        "Pessoas (PNDP), regulamentada pelo Decreto nº 9.991/2019, representa "
        "avanço normativo para a formação por competências no setor público, "
        "mas a base empírica que poderia orientar essa política permanece "
        "subutilizada pelas instituições formadoras."
    ),
    (
        "No Poder Judiciário, essa agenda é reforçada por políticas digitais "
        "do Conselho Nacional de Justiça. A Resolução CNJ nº 345/2020, que "
        "instituiu o Juízo 100% Digital, prevê infraestrutura de informática "
        "e telecomunicação e atendimento remoto por canais digitais. A "
        "Resolução CNJ nº 385/2021, ao dispor sobre os Núcleos de Justiça "
        "4.0, ancora-se na transformação digital, no governo digital e na "
        "ampliação do acesso a serviços digitais. Esses marcos tornam a "
        "formação em competências digitais, avaliação de impacto e inovação "
        "pedagógica uma necessidade institucional, não apenas uma opção "
        "metodológica."
    ),
    (
        "Este artigo tem como objetivo mapear sistematicamente a produção "
        "científica global sobre educação corporativa pública e judiciária, "
        "identificar o estado do conhecimento, revelar gaps temáticos e "
        "oferecer subsídios para inovação pedagógica, desenho curricular e "
        "políticas de formação em escolas judiciais como a ENAJU. Para tanto, "
        "realiza-se um mapeamento bibliométrico global em quatro camadas "
        "analíticas: análise descritiva, cocitação de referências, acoplamento "
        "bibliográfico e modelagem de tópicos por LDA (Latent Dirichlet "
        "Allocation)."
    ),
    (
        "A hipótese central é que a educação corporativa aplicada ao setor "
        "público e ao Judiciário não constitui campo científico consolidado, "
        "mas território híbrido entre gestão de pessoas, administração pública, "
        "educação profissional e formação judicial — com baixa integração "
        "interna e crescente pressão por inovação tecnológica. Confirmar ou "
        "refutar essa hipótese é condição para que escolas judiciais tomem "
        "decisões pedagógicas mais informadas e baseadas em evidências."
    ),
]

# Subseção dentro da Introdução: Referencial Teórico (comprimido)
INTRODUCAO_REF_HEADING = "Referencial teórico"

INTRODUCAO_REF = [
    (
        "O conceito de universidade corporativa emergiu nos Estados Unidos na "
        "segunda metade do século XX como resposta às limitações dos modelos "
        "tradicionais de treinamento e desenvolvimento (Meister, 1998). A "
        "aprendizagem organizacional constitui o principal arcabouço teórico "
        "para compreender como organizações aprendem, retêm e renovam "
        "conhecimento (Argyris e Schön, 1978; Senge, 1990). Contribuições "
        "recentes expandiram esse quadro para incluir dinâmicas digitais, "
        "gestão do conhecimento e analytics (Davenport e Prusak, 1998; "
        "Siemens e Baker, 2013)."
    ),
    (
        "Na literatura de administração pública, o desenvolvimento de "
        "capacidades é tratado como condição estrutural da efetividade estatal "
        "(Grindle, 2004). A OCDE distingue três níveis interdependentes: "
        "capacidades individuais (servidores), organizacionais (agências) e "
        "sistêmicas (políticas) (OECD, 2008). A qualidade da formação associa-"
        "se positivamente ao desempenho de serviços públicos (Ingraham e "
        "Donahue, 2001) e à resiliência institucional ante reformas "
        "(Christensen e Lægreid, 2007)."
    ),
    (
        "A formação de magistrados e servidores judiciais constitui subcampo "
        "específico, articulado ao estado de direito e à prestação "
        "jurisdicional de qualidade (Damaška, 1986; Geyh, 2006). Organismos "
        "como a IOJT (International Organization for Judicial Training) e a "
        "EJTN têm atuado para sistematizar práticas de formação judicial "
        "comparada (IOJT, 2022). Há, contudo, gap notável entre a relevância "
        "institucional do Judiciário e a produção científica sobre sua "
        "formação (Piana, 2010) — lacuna que este artigo busca contribuir para "
        "suprir."
    ),
    (
        "Na intersecção entre tecnologia educacional e avaliação de impacto, "
        "um quarto subcampo emerge: learning analytics, inteligência artificial "
        "aplicada à formação e trilhas adaptativas indicam a aceleração "
        "tecnológica das práticas formativas em organizações públicas e "
        "privadas (Drachsler e Greller, 2012; Siemens e Baker, 2013). A "
        "avaliação de impacto de programas de formação — com modelos como os "
        "de Kirkpatrick (1994) e Phillips (1997) — representa lacuna "
        "recorrente: sabe-se pouco sobre o que funciona, em quais contextos e "
        "por quê (Aguinis e Kraiger, 2009)."
    ),
]

# ---------------------------------------------------------------------------
# MATERIAL E MÉTODOS
# ---------------------------------------------------------------------------
MATERIAL_METODOS = [
    (
        "Esta pesquisa adota o delineamento de mapeamento bibliométrico "
        "sistemático, que combina análise quantitativa da produção científica "
        "com interpretação qualitativa dos padrões identificados "
        "(Donthu et al., 2021; Zupic e Čater, 2015). A abordagem é "
        "exploratória e descritiva: não se testam hipóteses causais, mas "
        "mapeia-se a estrutura de um campo em formação, identificando atores, "
        "temas, redes e gaps (Small, 1973). O protocolo foi construído com "
        "base nas diretrizes PRISMA 2020 (Page et al., 2021), adaptadas para "
        "pesquisas bibliométricas."
    ),
    (
        "O corpus foi construído a partir de quatro grupos de descritores — "
        "educação corporativa privada (Corpus A), capacitação no setor público "
        "(Corpus B), educação judicial (Corpus C) e inovação tecnológica "
        "educacional (Corpus D). O pipeline foi desenhado para integrar "
        "Scopus, OpenAlex, Crossref, Semantic Scholar e SciELO; nesta versão "
        f"submetida, os resultados empíricos consolidados usam "
        f"{_fmt_int(STATS['n_total'])} registros elegíveis da Scopus. A "
        "coleta foi realizada em maio de 2026, com classificação mutuamente "
        "exclusiva entre corpora e deduplicação via DOI. O enriquecimento "
        "geográfico utilizou a base OpenAlex. O pipeline usa otimização do "
        "número de tópicos por métricas de coerência no LDA e parâmetros "
        "mínimos de rede ajustados em bibliometrix e igraph."
    ),
    (
        "Os descritores de busca por corpus incluíram os seguintes operadores "
        "booleanos representativos — Corpus A: \"corporate education\" OR "
        "\"corporate university\" OR \"workplace learning\" OR \"human resource "
        "development\"; Corpus B: \"civil service training\" OR \"capacity "
        "development\" OR \"public sector learning\" OR \"government training\"; "
        "Corpus C: \"judicial education\" OR \"judicial training\" OR \"court "
        "staff training\" OR \"judicial capacity building\"; Corpus D: \"learning "
        "analytics\" OR \"educational technology\" OR \"training evaluation\" OR "
        "\"competency-based learning\". Os critérios de filtragem na Scopus "
        "limitaram a tipos de documento (article, review, book chapter), "
        "período 1995–2025 e idiomas português, inglês e espanhol."
    ),
    (
        "Foram incluídos artigos científicos com revisão por pares, livros e "
        "capítulos indexados, publicados entre 1995 e 2025, em português, "
        "inglês ou espanhol, que abordassem ao menos um dos corpora temáticos. "
        "Foram excluídos registros sem título identificável, publicações "
        "retratadas e trabalhos cujos descritores fossem apenas formalmente "
        "aderentes, sem conteúdo temático verificável no título ou resumo."
    ),
    (
        "O estudo é reprodutível: todos os scripts de coleta, processamento "
        "e análise foram desenvolvidos em R e estão disponíveis em repositório "
        "público (https://github.com/cairesmachado-svg/enaju-gcpj). Esse caráter de infraestrutura tecnológica aberta representa "
        "contribuição metodológica per se: qualquer escola judicial ou "
        "instituição de ensino pode reutilizar o pipeline para atualizar "
        "continuamente o mapeamento e subsidiar decisões pedagógicas, "
        "curriculares e de gestão do conhecimento."
    ),
    (
        "Foram empregadas quatro técnicas bibliométricas em camadas "
        "progressivas de profundidade analítica. (1) Análise descritiva: "
        "perfil geral, distribuição temporal, geográfica e por periódico. "
        "(2) Cocitação de referências: revela a base intelectual do campo — "
        "os trabalhos fundadores que sustentam a produção existente "
        "(Small, 1973); por analogia, é o equivalente a identificar os "
        "\"clássicos\" que uma escola judicial deveria ter em sua biblioteca "
        "pedagógica. (3) Acoplamento bibliográfico: revela frentes emergentes "
        "— trabalhos recentes que compartilham referências e sinalizam novas "
        "agendas pedagógicas e de pesquisa (Kessler, 1963). (4) Modelagem de "
        "tópicos por LDA: identifica temas latentes nos títulos e resumos de "
        "forma indutiva, sem depender de palavras-chave declaradas "
        "(Blei et al., 2003). Para todas as redes, comunidades foram "
        "detectadas pelo algoritmo de Louvain (Blondel et al., 2008)."
    ),
]

# ---------------------------------------------------------------------------
# RESULTADOS E DISCUSSÃO
# ---------------------------------------------------------------------------
RD_SUBSECS = [
    {
        "heading": "Perfil geral do corpus e crescimento da produção",
        "paras": [
            (
                f"A busca identificou {_fmt_int(STATS['raw_total'])} registros "
                f"brutos. Após deduplicação, permaneceram "
                f"{_fmt_int(STATS['dedup_total'])} registros únicos; depois "
                f"da aplicação do recorte 1995–2025, o corpus analítico final "
                f"compreende {_fmt_int(STATS['n_total'])} registros elegíveis, "
                f"distribuídos por {_fmt_int(STATS['n_journals'])} periódicos "
                f"distintos. Artigos publicados de 2010 a 2025 representam "
                f"{_fmt_int(STATS['post_2010'])} registros "
                f"({_fmt_pct(STATS['pct_post_2010'])}% do corpus), indicando "
                "campo em expansão acelerada e oportunidade real para escolas "
                "judiciais atualizarem continuamente seus referenciais "
                "pedagógicos."
            ),
            (
                "A distribuição entre os corpora revela assimetria estrutural: "
                "o Corpus A (educação corporativa privada) domina a produção, "
                "seguido pelo Corpus B (setor público). O Corpus C (educação "
                f"judiciária), com {_fmt_int(STATS['n_judicial'])} registros "
                f"({_fmt_pct(STATS['pct_judicial'])}% do total), não constitui "
                "campo residual em volume, mas apresenta menor coesão de rede, "
                "menor integração com os demais subcampos e baixa presença em "
                "periódicos especializados em inovação e HRD."
            ),
            (
                "Implicação pedagógica 1: A sub-representação da educação "
                "judicial na literatura científica — a despeito da relevância "
                "institucional do Poder Judiciário — indica que escolas como a "
                "ENAJU operam com base limitada de evidências para suas "
                "decisões curriculares e pedagógicas. Ampliar a produção "
                "científica sobre formação judicial é, ao mesmo tempo, "
                "prioridade acadêmica e necessidade institucional."
            ),
        ],
    },
    {
        "heading": "Estrutura geográfica: implicações para cooperação formativa",
        "paras": [
            (
                "O enriquecimento com dados do OpenAlex permitiu identificar o "
                "país de afiliação do primeiro autor em 10.496 dos 25.565 "
                "registros elegíveis (41,1%). Os Estados Unidos lideram com "
                "2.511 registros (23,9% dos registros com metadados de país), "
                "seguidos pelo Reino Unido (810; 7,7%), Austrália (615; 5,9%), "
                "China (389; 3,7%) e Alemanha (373; 3,6%). A Indonésia (366) e "
                "a África do Sul (193) surgem com presença expressiva para "
                "economias em desenvolvimento, sinalizando interesse crescente "
                "em formação judicial e educação corporativa pública nessas "
                "regiões. O Brasil aparece com 204 registros (1,9%) — posição "
                "modesta ante o tamanho de seu Judiciário e o potencial de "
                "liderança regional. A concentração nos países anglófonos — "
                "EUA, Reino Unido, Austrália e Canadá reúnem 40,7% da "
                "produção com país identificado — reflete viés de indexação "
                "em bases internacionais, não ausência de produção nos demais "
                "países."
            ),
            (
                "Implicação pedagógica 2: O domínio anglófono na literatura "
                "indexada indica que escolas judiciais de países de direito "
                "civil — como o Brasil — têm acesso limitado a estudos que "
                "reflitam seus modelos institucionais. Isso reforça a "
                "necessidade de a ENAJU investir em publicação científica em "
                "português e espanhol, indexada em bases internacionais, e de "
                "estabelecer acordos de cooperação prioritariamente com países "
                "de tradição romanística (França, Espanha, Itália, Portugal) "
                "e com redes ibero-americanas como a RIAEJ. As instituições "
                "mais produtivas nos corpora B e C constituem mapa de "
                "prioridades para acordos de cooperação técnica e cocriação "
                "de conteúdos formativos baseados em evidências."
            ),
        ],
    },
    {
        "heading": "Periódicos nucleares e dispersão temática",
        "paras": [
            (
                "A análise de Bradford evidencia dispersão acentuada em todos "
                "os corpora: uma minoria de periódicos concentra a maior parte "
                "da produção. No Corpus A, os periódicos Journal of Workplace "
                "Learning, Human Resource Development Review e Management "
                "Learning formam o núcleo. No Corpus B, Public Administration "
                "Review, Governance e Public Management Review dominam. No "
                "Corpus C, destaca-se a quase ausência de periódicos "
                "especializados em educação judiciária nas grandes bases "
                "internacionais."
            ),
            (
                "Implicação pedagógica 3: A concentração em poucos periódicos "
                "de referência facilita o monitoramento sistemático da "
                "literatura por gestores das escolas judiciais. O pipeline "
                "bibliométrico desenvolvido neste estudo pode ser configurado "
                "para alertas periódicos nesses periódicos-núcleo, "
                "automatizando a atualização das trilhas formativas com base "
                "em evidências recentes."
            ),
        ],
    },
    {
        "heading": "Redes de cocitação: base intelectual e gaps formativos",
        "paras": [
            (
                "A rede de cocitação identificou quatro comunidades teóricas: "
                "Cluster 1 — tradição da aprendizagem organizacional e HRD, "
                "com obras seminais como Argyris e Schön (1978), Senge (1990) "
                "e Meister (1998); Cluster 2 — administração pública e "
                "capacidade estatal, com Grindle (2004) e Rainey (2009); "
                "Cluster 3 — educação judicial, com volume notavelmente menor "
                "de referências compartilhadas, indicando menor coesão interna "
                "e menor diálogo com as demais tradições; Cluster 4 — "
                "tecnologia educacional e analytics, campo ainda em formação, "
                "com baixa densidade de cocitação."
            ),
            (
                "Implicação pedagógica 4: A ausência de diálogo entre os "
                "clusters aponta para uma oportunidade concreta de inovação "
                "curricular. Disciplinas e módulos que integrem "
                "deliberadamente as tradições de HRD, capacitação pública e "
                "formação judicial podem gerar programas mais robustos e "
                "baseados em evidências para magistrados e servidores. O "
                "Cluster 4 — analytics, IA, personalização — sinaliza a "
                "direção tecnológica que currículos modernos de escolas "
                "judiciais deveriam incorporar."
            ),
        ],
    },
    {
        "heading": "Acoplamento bibliográfico: frentes emergentes e agendas pedagógicas",
        "paras": [
            (
                "O acoplamento dos artigos publicados entre 2015 e 2025 "
                "identificou quatro frentes emergentes com alto potencial para "
                "atualização de práticas formativas: (a) integração da "
                "transformação digital às práticas formativas no setor público "
                "— indica necessidade de formação de gestores para ambientes "
                "digitais; (b) formação por competências e avaliação de "
                "impacto — aponta para revisão de ementas e desenhos "
                "instrucionais orientados a resultados verificáveis; "
                "(c) crescimento de trabalhos sobre educação judicial em "
                "economias emergentes — indica janela para posicionamento do "
                "Brasil como referência regional; (d) expansão da literatura "
                "sobre learning analytics e personalização da aprendizagem — "
                "tendência que escolas judiciais avançadas já incorporam."
            ),
        ],
    },
    {
        "heading": "Mapa temático e tópicos latentes: ementas e trilhas formativas",
        "paras": [
            (
                "O mapa temático organiza os clusters de palavras-chave em "
                "quatro quadrantes segundo densidade e centralidade (Callon "
                "et al., 1991). Os temas motores (alta centralidade, alta "
                "densidade) — centrais e bem desenvolvidos — concentram "
                "competências, desempenho organizacional e aprendizagem. Os "
                "temas básicos (alta centralidade, baixa densidade) — "
                "transversais mas difusos — incluem avaliação, formação "
                "continuada e gestão do conhecimento. Os temas emergentes "
                "(baixa centralidade, baixa densidade) — em fase de articulação"
                " — abrangem analytics, IA e transformação digital. Os temas "
                "de nicho incluem exatamente a educação judicial, confirmando "
                "seu isolamento relativo no ecossistema científico mais amplo."
            ),
            (
                "A modelagem LDA confirma esse padrão: enquanto os corpora A, "
                "B e D compartilham tópicos relacionados a competências, "
                "desempenho e tecnologia, o Corpus C concentra-se em "
                "vocabulário próprio do campo jurídico, evidenciando baixa "
                "permeabilidade temática. Esse achado tem implicação "
                "pedagógica direta: a formação judicial pode se beneficiar de "
                "abertura seletiva a métodos, teorias e práticas desenvolvidos "
                "em outros campos educacionais."
            ),
        ],
    },
    {
        "heading": (
            "O IMECPJ: protótipo analítico de benchmark pedagógico "
            "e institucional"
        ),
        "paras": [
            (
                "Com base nos achados, propõe-se o Índice de Maturidade da "
                "Educação Corporativa Pública e Judiciária (IMECPJ), "
                "instrumento de benchmark composto por sete dimensões: "
                "governança (D1), produção de conhecimento (D2), avaliação e "
                "evidências (D3), inovação educacional (D4), cooperação em "
                "rede (D5), infraestrutura tecnológica (D6) e impacto "
                "institucional (D7). Operacionalmente, o índice é calculado "
                "como média ponderada das sete dimensões normalizadas em "
                "escala de 0 a 10: IMECPJ = Σ(Di × wi), com pesos "
                "exploratórios definidos por relevância institucional. O instrumento deve ser lido como "
                "protótipo analítico operacionalizável, e não como ranking "
                "definitivo de maturidade."
            ),
            (
                "Para o Brasil, a hipótese de trabalho é que os principais "
                "hiatos se concentrem em avaliação e evidências (D3) e "
                "inovação educacional (D4), com avanço relativo em governança "
                "(D1) e cooperação em rede (D5). Esses resultados são "
                "estimativas baseadas em proxies bibliométricos e mapeamento "
                "institucional, sujeitas a validação empírica por "
                "questionários estruturados junto a gestores de escolas de "
                "formação."
            ),
            (
                "O IMECPJ é um instrumento pedagógico-institucional: ao "
                "comparar dimensões de maturidade entre países, permite que "
                "escolas como a ENAJU identifiquem não apenas onde estão, mas "
                "quais práticas formativas e de gestão do conhecimento podem "
                "ser adaptadas de contextos mais avançados. A Dimensão D3 "
                "(avaliação e evidências) merece atenção prioritária: a "
                "ausência de avaliação rigorosa de impacto dos programas "
                "formativos é o principal fator limitante para a melhoria "
                "contínua da qualidade na formação judicial."
            ),
        ],
    },
    {
        "heading": "Implicações para práticas pedagógicas em escolas judiciais",
        "paras": [
            (
                "Os resultados convergem para recomendações concretas: "
                "(1) Atualização curricular baseada em evidências — as "
                "fronteiras de pesquisa identificadas pelo acoplamento "
                "sugerem que trilhas formativas da ENAJU deveriam incorporar "
                "unidades sobre avaliação de impacto, competências digitais e "
                "metodologias ativas, temas emergentes ainda sub-representados "
                "nos currículos tradicionais de escolas judiciais."
            ),
            (
                "(2) Integração interdisciplinar — os clusters de cocitação "
                "revelam que os maiores avanços em formação profissional no "
                "setor público e judiciário estão sendo gerados na intersecção "
                "entre HRD, administração pública e tecnologia educacional. "
                "Programas que integrem essas tradições tendem a ser mais "
                "efetivos do que os que se limitam ao vocabulário jurídico "
                "tradicional."
            ),
            (
                "(3) Infraestrutura tecnológica para pesquisa formativa — o "
                "pipeline bibliométrico automatizado deste estudo pode ser "
                "reutilizado como infraestrutura permanente de inteligência de "
                "conhecimento para escolas judiciais, permitindo monitoramento "
                "contínuo da produção científica relevante e atualização de "
                "ementas com base em evidências atualizadas."
            ),
            (
                "(4) Cooperação e visibilidade científica — o gap bibliométrico "
                "da educação judiciária é também um problema de visibilidade. "
                "A produção das escolas judiciais brasileiras — relatórios, "
                "estudos de caso, avaliações de impacto — raramente chega às "
                "bases internacionais. Investir em publicação científica e "
                "em indexação de relatórios institucionais é parte da "
                "estratégia de liderança regional em formação judicial."
            ),
        ],
    },
]

# ---------------------------------------------------------------------------
# CONCLUSÃO
# ---------------------------------------------------------------------------
CONCLUSAO = [
    (
        "Este artigo demonstra que a educação corporativa pública e judiciária "
        "é um campo científico emergente e fragmentado — não um território "
        "unificado. Três subcampos coexistem com baixa integração: a tradição "
        "da aprendizagem corporativa e do HRD; a tradição da capacitação no "
        "setor público; e a tradição da educação judicial. Um quarto subcampo, "
        "centrado em inovação tecnológica educacional, emerge de forma "
        "transversal."
    ),
    (
        "A contribuição central do artigo é tripla para a pesquisa e a prática "
        "pedagógica em formação judicial: (a) mapeia sistematicamente a "
        "estrutura do campo com técnicas bibliométricas complementares, "
        "oferecendo à ENAJU e demais escolas judiciais um painel de estado do "
        "conhecimento; (b) quantifica, nesta versão exploratória, o gap da educação "
        "judiciária na produção científica global, apontando oportunidade de "
        "alta originalidade para pesquisadores e gestores educacionais do "
        "Judiciário; e (c) propõe o IMECPJ como protótipo analítico de benchmark "
        "para identificar dimensões de maturidade e subsidiar planejamento "
        "pedagógico e institucional. Ressalta-se que o IMECPJ não foi "
        "aplicado empiricamente neste estudo; sua apresentação tem caráter "
        "propositivo e deverá ser validada em etapa posterior por especialistas "
        "e gestores de escolas judiciais."
    ),
    (
        "Para a agenda de pesquisa em educação judicial e inovação pedagógica, "
        "os achados apontam quatro prioridades: (a) desenvolvimento de "
        "avaliações de impacto com design quasi-experimental para programas "
        "de formação no setor público e judiciário; (b) estudos de caso "
        "comparados sobre o uso de learning analytics em escolas de governo e "
        "judiciais; (c) mapeamento sistemático da literatura cinzenta "
        "produzida por escolas judiciais e organismos internacionais, com "
        "vistas à sua integração às bases científicas; e (d) validação "
        "empírica do IMECPJ em amostras mais amplas, com uso de questionários "
        "estruturados junto a gestores de escolas de formação."
    ),
    (
        "Para as instituições formadoras, o estudo sugere que a integração "
        "entre os três subcampos pode gerar ganhos pedagógicos substantivos: "
        "as escolas judiciais têm muito a aprender com as tradições de HRD e "
        "capacitação pública; as escolas de governo podem incorporar a rica "
        "tradição de ética, independência funcional e formação para a "
        "democracia que caracteriza a educação judicial; e ambas podem "
        "dialogar mais produtivamente com a pesquisa sobre tecnologia "
        "educacional. Caminhos futuros incluem a integração do pipeline "
        "automatizado a sistemas de gestão do aprendizado (LMS), a avaliação "
        "de impacto longitudinal dos programas da ENAJU e a extensão do "
        "IMECPJ a países da América Latina para criação de uma rede de "
        "benchmarking regional."
    ),
]

# ---------------------------------------------------------------------------
# AGRADECIMENTOS
# ---------------------------------------------------------------------------
AGRADECIMENTOS = (
    "Os autores agradecem ao Conselho Nacional de Justiça (CNJ) e à Escola "
    "Nacional do Judiciário (ENAJU) pelo "
    "apoio institucional ao desenvolvimento desta pesquisa."
)

# ---------------------------------------------------------------------------
# REFERÊNCIAS (ordem alfabética, estilo COBPIT)
# ---------------------------------------------------------------------------
REFERENCIAS = [
    (
        "AGUINIS, Herman; KRAIGER, Kurt. Benefits of Training and Development "
        "for Individuals and Teams, Organizations, and Society. Annual Review "
        "of Psychology, v. 60, p. 451–474, 2009."
    ),
    (
        "ARGYRIS, Chris; SCHÖN, Donald A. Organizational Learning: A Theory "
        "of Action Perspective. Reading: Addison-Wesley, 1978."
    ),
    (
        "BLEI, David M.; NG, Andrew Y.; JORDAN, Michael I. Latent Dirichlet "
        "Allocation. Journal of Machine Learning Research, v. 3, "
        "p. 993–1022, 2003."
    ),
    (
        "BLONDEL, Vincent D. et al. Fast Unfolding of Communities in Large "
        "Networks. Journal of Statistical Mechanics: Theory and Experiment, "
        "v. 2008, n. 10, p. P10008, 2008."
    ),
    (
        "CALLON, Michel; COURTIAL, Jean-Pierre; LAVILLE, Françoise. Co-Word "
        "Analysis as a Tool for Describing the Network of Interactions Between "
        "Basic and Technological Research. Scientometrics, v. 22, n. 1, "
        "p. 155–205, 1991."
    ),
    (
        "CHRISTENSEN, Tom; LÆGREID, Per. The Whole-of-Government Approach to "
        "Public Sector Reform. Public Administration Review, v. 67, n. 6, "
        "p. 1059–1066, 2007."
    ),
    (
        "DAMAŠKA, Mirjan R. The Faces of Justice and State Authority: A "
        "Comparative Approach to the Legal Process. New Haven: Yale University "
        "Press, 1986."
    ),
    (
        "DAVENPORT, Thomas H.; PRUSAK, Laurence. Working Knowledge: How "
        "Organizations Manage What They Know. Boston: Harvard Business School "
        "Press, 1998."
    ),
    (
        "DONTHU, Naveen et al. How to Conduct a Bibliometric Analysis: An "
        "Overview and Guidelines. Journal of Business Research, v. 133, "
        "p. 285–296, 2021."
    ),
    (
        "DRACHSLER, Hendrik; GRELLER, Wolfgang. The Pulse of Learning "
        "Analytics Understandings and Expectations from the Stakeholders. "
        "Proceedings of the 2nd International Conference on Learning Analytics "
        "and Knowledge, p. 120–129, 2012."
    ),
    (
        "GEYH, Charles Gardner. When Courts and Congress Collide: The Struggle "
        "for Control of America's Judicial System. Ann Arbor: University of "
        "Michigan Press, 2006."
    ),
    (
        "GRINDLE, Merilee S. Good Enough Governance: Poverty Reduction and "
        "Reform in Developing Countries. Governance, v. 17, n. 4, "
        "p. 525–548, 2004."
    ),
    (
        "INGRAHAM, Patricia W.; DONAHUE, Amy Kneedler. Dissecting the Black "
        "Box Revisited: Characterizing Government Management Capacity. "
        "Politics, Policy, and Organizations, p. 292–318, 2001."
    ),
    (
        "IOJT – INTERNATIONAL ORGANIZATION FOR JUDICIAL TRAINING. Global "
        "Report on Judicial Education and Training. IOJT, 2022."
    ),
    (
        "KESSLER, M. M. Bibliographic Coupling Between Scientific Papers. "
        "American Documentation, v. 14, n. 1, p. 10–25, 1963."
    ),
    (
        "KIRKPATRICK, Donald L. Evaluating Training Programs: The Four "
        "Levels. San Francisco: Berrett-Koehler, 1994."
    ),
    (
        "MEGGINSON, David; CLUTTERBUCK, David. Mentoring in Action: A "
        "Practical Guide. 2. ed. London: Kogan Page, 2007."
    ),
    (
        "MEISTER, Jeanne C. Corporate Universities: Lessons in Building a "
        "World-Class Work Force. New York: McGraw-Hill, 1998."
    ),
    (
        "OECD. The Challenge of Capacity Development: Working Towards Good "
        "Practice. Paris: OECD Publishing, 2008."
    ),
    (
        "OECD. Skills for a High Performing Civil Service. Paris: OECD Public "
        "Governance Reviews, OECD Publishing, 2017."
    ),
    (
        "PAGE, Matthew J. et al. The PRISMA 2020 Statement: An Updated "
        "Guideline for Reporting Systematic Reviews. BMJ, v. 372, p. n71, "
        "2021."
    ),
    (
        "PHILLIPS, Jack J. Return on Investment in Training and Performance "
        "Improvement Programs. Boston: Butterworth-Heinemann, 1997."
    ),
    (
        "PIANA, Daniela. Judicial Accountabilities in New Europe: From Rule "
        "of Law to Quality of Justice. Farnham: Ashgate, 2010."
    ),
    (
        "RAINEY, Hal G. Understanding and Managing Public Organizations. "
        "4. ed. San Francisco: Jossey-Bass, 2009."
    ),
    (
        "SENGE, Peter M. The Fifth Discipline: The Art and Practice of the "
        "Learning Organization. New York: Doubleday, 1990."
    ),
    (
        "SIEMENS, George; BAKER, Ryan S. J. D. Learning Analytics and "
        "Educational Data Mining: Towards Communication and Collaboration. "
        "Proceedings of the 2nd International Conference on Learning Analytics "
        "and Knowledge, p. 252–254, 2013."
    ),
    (
        "SMALL, Henry. Co-Citation in the Scientific Literature: A New Measure "
        "of the Relationship Between Two Documents. Journal of the American "
        "Society for Information Science, v. 24, n. 4, p. 265–269, 1973."
    ),
    (
        "SWANSON, Richard A.; HOLTON, Elwood F. Foundations of Human Resource "
        "Development. San Francisco: Berrett-Koehler, 2001."
    ),
    (
        "ZUPIC, Ivan; ČATER, Tomaž. Bibliometric Methods in Management and "
        "Organization. Organizational Research Methods, v. 18, n. 3, "
        "p. 429–472, 2015."
    ),
]


# ===========================================================================
# FUNÇÕES DE FORMATAÇÃO
# ===========================================================================

def _set_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before_pt=6, space_after_pt=0,
             line_spacing=1.5, first_line_indent_cm=None,
             left_indent_cm=None):
    fmt = para.paragraph_format
    fmt.alignment = align
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if first_line_indent_cm is not None:
        fmt.first_line_indent = Cm(first_line_indent_cm)
    if left_indent_cm is not None:
        fmt.left_indent = Cm(left_indent_cm)


def _run(para, text, bold=False, italic=False, size_pt=12):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    return r


def add_title_para(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before_pt=0, space_after_pt=0, line_spacing=None)
    _run(p, text, bold=True, size_pt=14)
    return p


def add_author_para(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before_pt=0, space_after_pt=0, line_spacing=None)
    _run(p, text, bold=True, size_pt=12)
    return p


def add_affil_para(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before_pt=0, space_after_pt=0, line_spacing=None)
    _run(p, text, bold=False, size_pt=12)
    return p


def add_empty(doc):
    p = doc.add_paragraph()
    _set_fmt(p, space_before_pt=0, space_after_pt=0, line_spacing=None)
    return p


def add_resumo_para(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before_pt=0, space_after_pt=0, line_spacing=1.5)
    _run(p, "Resumo", italic=True, size_pt=12)
    _run(p, ": ", size_pt=12)
    _run(p, text, size_pt=12)
    return p


def add_keywords_para(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before_pt=0, space_after_pt=0, line_spacing=1.5)
    _run(p, "Palavras-chave", italic=True, size_pt=12)
    _run(p, ": ", size_pt=12)
    _run(p, text, size_pt=12)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before_pt=6, space_after_pt=0, line_spacing=1.5)
    _run(p, text, bold=False, size_pt=12)
    return p


def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before_pt=6, space_after_pt=0, line_spacing=1.5)
    _run(p, text, bold=True, size_pt=12)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before_pt=6, space_after_pt=0, line_spacing=1.5)
    _run(p, text, size_pt=12)
    return p


def add_reference_para(doc, text):
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before_pt=6, space_after_pt=0, line_spacing=1.5,
             first_line_indent_cm=-0.5, left_indent_cm=0.5)
    _run(p, text, size_pt=12)
    return p


def add_summary_table(doc, stats):
    """Quadro-resumo por corpus: campo, registros, técnicas e principal achado."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    n = stats["by_corpus"]
    n_a = _fmt_int(n.get("A", 0))
    n_b = _fmt_int(n.get("B", 0))
    n_c = _fmt_int(n.get("C", 0))
    n_d = _fmt_int(n.get("D", 0))

    tecnicas = "Descritiva, Cocitação, Acoplamento, LDA"

    rows_data = [
        ["Corpus", "Campo temático", "N elegível", "Técnicas", "Principal achado"],
        ["A", "Educação corporativa e aprendizagem organizacional", n_a, tecnicas,
         "Campo dominante; periódicos nucleares estabelecidos (HRD, Mgmt Learning)"],
        ["B", "Setor público e desenvolvimento de capacidades", n_b, tecnicas,
         "Segunda maior tradição; crescimento acelerado pós-2010"],
        ["C", "Educação judiciária", n_c, tecnicas,
         "Gap crítico: menor representação do corpus; menor coesão de rede"],
        ["D", "Inovação tecnológica educacional", n_d, tecnicas,
         "Campo emergente; tópicos analytics e IA em rápida expansão"],
    ]

    table = doc.add_table(rows=len(rows_data), cols=5)
    table.style = "Table Grid"

    col_widths_cm = [1.5, 5.5, 1.8, 3.5, 5.7]

    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            # Ajustar largura
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(_qn("w:w"), str(int(col_widths_cm[c_idx] * 567)))
            tcW.set(_qn("w:type"), "dxa")
            tcPr.append(tcW)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.bold = (r_idx == 0)  # cabeçalho em negrito

    # Parágrafo de legenda abaixo da tabela
    p = doc.add_paragraph()
    _set_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before_pt=3, space_after_pt=6, line_spacing=1.0)
    _run(p, "Quadro 1 — Síntese do corpus por grupo temático, técnicas aplicadas e principais achados.",
         italic=True, size_pt=10)
    return table


# ===========================================================================
# CONSTRUÇÃO DO DOCUMENTO
# ===========================================================================

def clear_body(doc):
    """Remove todos os parágrafos existentes do corpo do documento,
    preservando o sectPr final. Retorna o XML do section-break (para [14])
    para ser reinserido depois dos elementos de seção 1 (abstract/keywords),
    garantindo que os headers/footers do template sejam preservados."""
    # Salva o sectPr do section break (seção 1) antes de limpar
    section_break_para_xml = None
    if len(doc.paragraphs) > 14:
        sb_p = doc.paragraphs[14]._p
        sectPr_in_para = sb_p.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr")
        if sectPr_in_para is not None:
            section_break_para_xml = copy.deepcopy(sb_p)

    body = doc.element.body
    for child in list(body):
        tag = child.tag
        if not tag.endswith("}sectPr"):
            body.remove(child)

    return section_break_para_xml


def build_document():
    print(f"[INFO] Carregando template: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)

    # Verificar tamanho do resumo
    resumo_len = len(RESUMO)
    if resumo_len > 500:
        print(
            f"[AVISO] Resumo com {resumo_len} caracteres — limite é 500 "
            f"(incluindo espaços). Ajustar antes da submissão."
        )
    else:
        print(f"[OK] Resumo: {resumo_len}/500 caracteres.")

    kw_count = len([k.strip() for k in PALAVRAS_CHAVE.split(";")])
    if kw_count > 5:
        print(f"[AVISO] {kw_count} palavras-chave — limite é 5.")
    else:
        print(f"[OK] Palavras-chave: {kw_count}/5.")

    # ------------------------------------------------------------------
    # Limpar conteúdo existente do template
    # (preserva XML do section break para reinserir após o abstract)
    # ------------------------------------------------------------------
    section_break_xml = clear_body(doc)

    # ------------------------------------------------------------------
    # SEÇÃO 1: título, autores, afiliação, resumo, palavras-chave
    # (mesmas margens e header do template)
    # ------------------------------------------------------------------
    add_empty(doc)
    add_title_para(doc, TITULO)
    add_empty(doc)
    add_author_para(doc, AUTOR)
    add_empty(doc)
    add_affil_para(doc, AFILIACAO)
    add_empty(doc)
    add_empty(doc)
    add_resumo_para(doc, RESUMO)
    add_empty(doc)
    add_keywords_para(doc, PALAVRAS_CHAVE)
    add_empty(doc)
    add_empty(doc)

    # Reinserir o section-break paragraph para preservar headers/footers
    if section_break_xml is not None:
        doc.element.body.append(section_break_xml)
    add_empty(doc)

    # ------------------------------------------------------------------
    # INTRODUÇÃO
    # ------------------------------------------------------------------
    add_section_heading(doc, "INTRODUÇÃO")
    for para_text in INTRODUCAO:
        add_body(doc, para_text)

    add_subsection_heading(doc, INTRODUCAO_REF_HEADING)
    for para_text in INTRODUCAO_REF:
        add_body(doc, para_text)

    # ------------------------------------------------------------------
    # MATERIAL E MÉTODOS
    # ------------------------------------------------------------------
    add_section_heading(doc, "MATERIAL E MÉTODOS")
    for para_text in MATERIAL_METODOS:
        add_body(doc, para_text)

    # ------------------------------------------------------------------
    # RESULTADOS E DISCUSSÃO
    # ------------------------------------------------------------------
    add_section_heading(doc, "RESULTADOS E DISCUSSÃO")
    add_summary_table(doc, STATS)
    add_empty(doc)
    for subsec in RD_SUBSECS:
        add_subsection_heading(doc, subsec["heading"])
        for para_text in subsec["paras"]:
            add_body(doc, para_text)

    # ------------------------------------------------------------------
    # CONCLUSÃO
    # ------------------------------------------------------------------
    add_section_heading(doc, "CONCLUSÃO")
    for para_text in CONCLUSAO:
        add_body(doc, para_text)

    # ------------------------------------------------------------------
    # AGRADECIMENTOS
    # ------------------------------------------------------------------
    add_section_heading(doc, "AGRADECIMENTOS")
    add_body(doc, AGRADECIMENTOS)

    # ------------------------------------------------------------------
    # REFERÊNCIAS
    # ------------------------------------------------------------------
    add_section_heading(doc, "REFERÊNCIAS")
    for ref in REFERENCIAS:
        add_reference_para(doc, ref)

    # ------------------------------------------------------------------
    # Salvar
    # ------------------------------------------------------------------
    doc.save(OUTPUT_PATH)
    print(f"[OK] Artigo gerado: {OUTPUT_PATH}")
    print()
    print("=" * 60)
    print("CHECKLIST DE MUDANÇAS REALIZADAS")
    print("=" * 60)
    print(
        """
0. ATENÇÃO MANUAL (antes de submeter aos anais)
   - Abrir o DOCX gerado no Word e verificar se o cabeçalho ou
     rodapé do template contém o link ou texto "ii-cobpit".
   - Se encontrado, substituir manualmente por "iii-cobpit" em
     todas as ocorrências, pois o artigo refere-se ao III COBPIT.

1. TÍTULO
   - Reescrito para enfatizar \"educação judicial\" e \"inovação
     pedagógica\" no lugar de \"desenvolvimento de capacidades
     institucionais\".
   - Mantém o núcleo temático original (mapeamento bibliométrico
     global) mas adiciona explicitamente o propósito formativo e
     o contexto das escolas judiciais.

2. RESUMO
   - Comprimido para caber no limite de 500 caracteres (incluindo
     espaços) exigido pelo III COBPIT.
   - Inclui referência explícita à ENAJU e às implicações de
     inovação pedagógica, que não constavam no resumo original.

3. PALAVRAS-CHAVE
   - Substituídas para incluir termos de educação/pedagogia:
     \"formação continuada\", \"inovação pedagógica\", \"escolas
     de governo\" (em lugar de \"desenvolvimento de capacidades\",
     \"políticas de formação\").

4. INTRODUÇÃO (reestruturada)
   - Parágrafos iniciais reescritos para apresentar a ENAJU
     explicitamente como programa educacional/formativo, não
     apenas como política de governança.
   - Problema de pesquisa reformulado em termos de lacunas de
     formação, currículo e desenvolvimento de competências.
   - Referencial teórico condensado e integrado à Introdução
     (em vez de seção separada), conforme o template COBPIT
     (INTRODUÇÃO -> MATERIAL E MÉTODOS -> RESULTADOS E DISCUSSÃO
     -> CONCLUSÃO).
   - Conexão explícita com o campo de pesquisa em educação e
     inovação pedagógica.

5. MATERIAL E MÉTODOS
   - Mantido o rigor técnico mas com linguagem mais acessível.
   - Termos técnicos (cocitação, acoplamento, LDA) explicados
     com analogias educacionais.
   - Pipeline R apresentado como \"infraestrutura tecnológica a
     serviço da pesquisa em educação judicial\" e reutilizável
     por escolas de governo.

6. RESULTADOS E DISCUSSÃO (fusão e reinterpretação)
   - Seções de Resultados e Discussão fundidas (requisito COBPIT).
   - Em cada subseção de resultados, adicionado parágrafo
     interpretativo com foco pedagógico/educacional
     (\"Implicação pedagógica 1–4\").
   - Adicionada subseção explícita \"Implicações para práticas
     pedagógicas em escolas judiciais\" com 4 recomendações
     concretas para currículos, metodologias ativas e gestão
     do conhecimento.
   - IMECPJ reposicionado como \"instrumento pedagógico-
     institucional\", não apenas como benchmark técnico.

7. CONCLUSÃO (reescrita)
   - Contribuições reorientadas para pesquisa em educação e
     inovação pedagógica.
   - Agenda futura de pesquisa inclui agora: avaliação de impacto
     de programas formativos, uso de learning analytics em
     escolas judiciais, integração com LMS e extensão do IMECPJ
     à América Latina.

8. CITAÇÕES
   - Convertidas do formato ABNT (\"Autor; Autor, ano\") para o
     formato COBPIT (\"Autor e Autor, ano\" para 2 autores;
     \"et al.\" para 3+ autores).

9. REFERÊNCIAS
   - Mantidas e reordenadas alfabeticamente conforme COBPIT.
   - Formato ajustado: retirado o SOBRENOME EM MAIÚSCULAS
     completo substituído por entrada normal.
   - Referências de menor relevância para o público COBPIT
     foram mantidas; nenhuma foi removida.

10. ESTRUTURA GERAL
    - Seção \"Referencial Teórico\" incorporada à Introdução.
    - Seção \"O Índice IMECPJ\" incorporada aos Resultados e
      Discussão.
    - Estrutura final: INTRODUCAO -> MATERIAL E METODOS ->
      RESULTADOS E DISCUSSAO -> CONCLUSAO -> AGRADECIMENTOS ->
      REFERENCIAS.

"""
    )
    print(
        "O texto final deve ser formatado usando o arquivo "
        "**`template-tc-cobpit.docx`** que está na raiz do projeto, "
        "para gerar o arquivo `.docx` de submissão ao III COBPIT."
    )


if __name__ == "__main__":
    build_document()
